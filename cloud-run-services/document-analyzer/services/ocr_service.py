"""
LicitaReview - Serviço de OCR Avançado com Google Vision API
🚀 v2.0.0 - Integração Completa

Serviço responsável por extrair texto de documentos diversos,
incluindo PDFs, imagens e documentos Word com recursos avançados:
- OCR via Google Vision API
- Extração de tabelas
- Detecção de layout de documento
- Extração de formulários
- Análise de estrutura de documento
"""

import io
import os
import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime

# Google Cloud Vision
from google.cloud import vision
from google.cloud.vision_v1 import types

# PDF Processing
try:
    import PyPDF2
    from pdf2image import convert_from_bytes
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PyPDF2 ou pdf2image não disponível - OCR de PDF limitado")

# Word Processing
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx não disponível - processamento DOCX limitado")

# Image Processing
try:
    from PIL import Image
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False
    logging.warning("Pillow não disponível - processamento de imagem limitado")

logger = logging.getLogger(__name__)


@dataclass
class TableCell:
    """Célula de tabela extraída."""
    row: int
    col: int
    text: str
    confidence: float
    rowspan: int = 1
    colspan: int = 1


@dataclass
class ExtractedTable:
    """Tabela extraída do documento."""
    rows: int
    cols: int
    cells: List[TableCell]
    confidence: float
    position: Dict[str, int]  # {x, y, width, height}


@dataclass
class LayoutBlock:
    """Bloco de layout detectado."""
    type: str  # 'text', 'heading', 'list', 'table', 'image'
    text: str
    confidence: float
    position: Dict[str, int]
    metadata: Dict[str, Any]


@dataclass
class FormField:
    """Campo de formulário extraído."""
    field_name: str
    field_value: str
    field_type: str  # 'text', 'checkbox', 'signature', etc.
    confidence: float
    position: Dict[str, int]


@dataclass
class OCRResult:
    """Resultado completo da extração OCR."""
    text: str
    confidence: float
    language: str
    tables: List[ExtractedTable]
    layout_blocks: List[LayoutBlock]
    form_fields: List[FormField]
    metadata: Dict[str, Any]
    processing_time: float
    method: str  # 'vision_api', 'pypdf', 'docx', 'fallback'


class OCRService:
    """
    Serviço de extração de texto avançado com Google Vision API.

    Recursos:
    - OCR via Google Cloud Vision API
    - Extração de tabelas
    - Detecção de layout
    - Extração de formulários
    - Suporte a PDF, DOCX, imagens
    """

    def __init__(self):
        """Inicializa o serviço OCR."""
        self.logger = logging.getLogger(self.__class__.__name__)
        self.vision_client: Optional[vision.ImageAnnotatorClient] = None
        self.is_initialized = False
        self.use_vision_api = False

        # Estatísticas
        self.stats = {
            'total_extractions': 0,
            'vision_api_calls': 0,
            'fallback_calls': 0,
            'errors': 0
        }

    def initialize(self):
        """Inicializa o serviço de OCR com Google Vision API."""
        if self.is_initialized:
            return

        self.logger.info("🚀 Initializing OCRService v2.0.0")

        # Tentar inicializar Google Vision API
        try:
            self.vision_client = vision.ImageAnnotatorClient()
            self.use_vision_api = True
            self.logger.info("✅ Google Vision API inicializada com sucesso")
        except Exception as e:
            self.logger.warning(f"⚠️  Google Vision API não disponível: {e}")
            self.logger.info("📝 Usando método fallback para OCR")
            self.use_vision_api = False

        self.is_initialized = True
        self.logger.info(f"✅ OCRService initialized (Vision API: {self.use_vision_api})")

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Extrai texto simples do arquivo (compatibilidade).

        Args:
            file_content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo

        Returns:
            Texto extraído
        """
        result = self.extract_full(file_content, filename)
        return result.text

    def extract_full(
        self,
        file_content: bytes,
        filename: str,
        extract_tables: bool = True,
        detect_layout: bool = True,
        extract_forms: bool = True
    ) -> OCRResult:
        """
        Extração completa com recursos avançados.

        Args:
            file_content: Conteúdo do arquivo
            filename: Nome do arquivo
            extract_tables: Extrair tabelas
            detect_layout: Detectar layout
            extract_forms: Extrair campos de formulário

        Returns:
            OCRResult completo com todas as informações
        """
        start_time = datetime.now()
        self.stats['total_extractions'] += 1

        if not self.is_initialized:
            self.initialize()

        self.logger.info(f"📄 Extracting text from: {filename}")

        # Detectar tipo de arquivo
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''

        try:
            # Escolher método de extração
            if file_ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                result = self._extract_from_image(
                    file_content, filename,
                    extract_tables, detect_layout, extract_forms
                )
            elif file_ext == 'pdf':
                result = self._extract_from_pdf(
                    file_content, filename,
                    extract_tables, detect_layout, extract_forms
                )
            elif file_ext in ['doc', 'docx']:
                result = self._extract_from_docx(file_content, filename)
            else:
                # Tentar como texto puro
                text = file_content.decode('utf-8', errors='ignore')
                result = OCRResult(
                    text=text,
                    confidence=1.0,
                    language='pt',
                    tables=[],
                    layout_blocks=[],
                    form_fields=[],
                    metadata={'file_type': 'text'},
                    processing_time=0,
                    method='plain_text'
                )

            processing_time = (datetime.now() - start_time).total_seconds()
            result.processing_time = processing_time

            self.logger.info(
                f"✅ Extraction complete: {len(result.text)} chars, "
                f"{len(result.tables)} tables, "
                f"{len(result.layout_blocks)} blocks in {processing_time:.2f}s"
            )

            return result

        except Exception as e:
            self.stats['errors'] += 1
            self.logger.error(f"❌ Extraction failed for {filename}: {e}")
            raise

    def _extract_from_image(
        self,
        image_content: bytes,
        filename: str,
        extract_tables: bool,
        detect_layout: bool,
        extract_forms: bool
    ) -> OCRResult:
        """Extrai texto de imagem usando Google Vision API."""
        if self.use_vision_api and self.vision_client:
            return self._extract_with_vision_api(
                image_content, filename,
                extract_tables, detect_layout, extract_forms
            )
        else:
            return self._extract_image_fallback(image_content, filename)

    def _extract_with_vision_api(
        self,
        image_content: bytes,
        filename: str,
        extract_tables: bool,
        detect_layout: bool,
        extract_forms: bool
    ) -> OCRResult:
        """Extração usando Google Vision API com recursos avançados."""
        self.stats['vision_api_calls'] += 1

        image = types.Image(content=image_content)

        # 1. Text Detection (OCR básico)
        text_response = self.vision_client.text_detection(image=image)
        texts = text_response.text_annotations

        full_text = texts[0].description if texts else ""
        confidence = texts[0].confidence if texts and hasattr(texts[0], 'confidence') else 0.95

        # Detectar idioma
        language = 'pt'  # Default português
        if texts and len(texts) > 0:
            for text in texts[:5]:  # Analisar primeiros textos
                if hasattr(text, 'locale'):
                    language = text.locale
                    break

        tables = []
        layout_blocks = []
        form_fields = []

        # 2. Document Text Detection (mais detalhado)
        if detect_layout:
            doc_response = self.vision_client.document_text_detection(image=image)
            layout_blocks = self._parse_document_layout(doc_response)

        # 3. Detecção de tabelas (se solicitado)
        if extract_tables:
            tables = self._extract_tables_from_layout(layout_blocks)

        # 4. Detecção de formulários (se solicitado)
        if extract_forms:
            form_fields = self._extract_form_fields(layout_blocks, full_text)

        return OCRResult(
            text=full_text,
            confidence=confidence,
            language=language,
            tables=tables,
            layout_blocks=layout_blocks,
            form_fields=form_fields,
            metadata={
                'file_type': 'image',
                'filename': filename,
                'vision_api': True
            },
            processing_time=0,  # Será preenchido depois
            method='vision_api'
        )

    def _parse_document_layout(self, doc_response) -> List[LayoutBlock]:
        """Analisa layout do documento a partir da resposta do Vision API."""
        layout_blocks = []

        if not doc_response.full_text_annotation:
            return layout_blocks

        for page in doc_response.full_text_annotation.pages:
            for block in page.blocks:
                # Determinar tipo de bloco
                block_type = self._determine_block_type(block)

                # Extrair texto do bloco
                block_text = ""
                for paragraph in block.paragraphs:
                    for word in paragraph.words:
                        word_text = ''.join([symbol.text for symbol in word.symbols])
                        block_text += word_text + " "

                # Calcular confiança
                confidence = block.confidence if hasattr(block, 'confidence') else 0.9

                # Extrair posição
                vertices = block.bounding_box.vertices
                position = {
                    'x': vertices[0].x,
                    'y': vertices[0].y,
                    'width': vertices[2].x - vertices[0].x,
                    'height': vertices[2].y - vertices[0].y
                }

                layout_blocks.append(LayoutBlock(
                    type=block_type,
                    text=block_text.strip(),
                    confidence=confidence,
                    position=position,
                    metadata={}
                ))

        return layout_blocks

    def _determine_block_type(self, block) -> str:
        """Determina o tipo de bloco baseado em características."""
        # Análise simples baseada no conteúdo
        # Em produção, usaria ML mais sofisticado

        if hasattr(block, 'block_type'):
            return str(block.block_type).lower()

        # Heurísticas simples
        return 'text'

    def _extract_tables_from_layout(self, layout_blocks: List[LayoutBlock]) -> List[ExtractedTable]:
        """Extrai tabelas a partir dos blocos de layout."""
        tables = []

        # Agrupar blocos que parecem formar tabelas
        # Implementação simplificada - em produção usaria algoritmo mais sofisticado

        table_blocks = [b for b in layout_blocks if 'tabela' in b.text.lower() or self._looks_like_table(b.text)]

        for block in table_blocks:
            # Analisar estrutura de tabela
            cells = self._parse_table_cells(block.text)

            if cells:
                # Determinar dimensões
                max_row = max(c.row for c in cells) + 1
                max_col = max(c.col for c in cells) + 1

                tables.append(ExtractedTable(
                    rows=max_row,
                    cols=max_col,
                    cells=cells,
                    confidence=block.confidence,
                    position=block.position
                ))

        return tables

    def _looks_like_table(self, text: str) -> bool:
        """Verifica se o texto parece uma tabela."""
        # Heurísticas simples
        lines = text.split('\n')
        if len(lines) < 2:
            return False

        # Verificar se há alinhamento consistente
        separators = ['|', '\t', '  ']
        for sep in separators:
            if all(sep in line for line in lines[:3]):
                return True

        return False

    def _parse_table_cells(self, table_text: str) -> List[TableCell]:
        """Analisa texto de tabela e extrai células."""
        cells = []
        lines = table_text.split('\n')

        for row_idx, line in enumerate(lines):
            # Dividir por separadores comuns
            parts = [p.strip() for p in line.split('|') if p.strip()]

            for col_idx, part in enumerate(parts):
                cells.append(TableCell(
                    row=row_idx,
                    col=col_idx,
                    text=part,
                    confidence=0.85,
                    rowspan=1,
                    colspan=1
                ))

        return cells

    def _extract_form_fields(self, layout_blocks: List[LayoutBlock], full_text: str) -> List[FormField]:
        """Extrai campos de formulário do documento."""
        form_fields = []

        # Padrões comuns de formulários em licitações
        field_patterns = {
            'numero_edital': r'edital\s*(?:n[º°]?|número)?\s*:?\s*(\d+/\d+)',
            'objeto': r'objeto\s*:?\s*(.+?)(?:\n|$)',
            'valor_estimado': r'valor\s*estimado\s*:?\s*r?\$?\s*([\d.,]+)',
            'prazo': r'prazo\s*(?:de)?\s*(?:entrega|execução)\s*:?\s*(.+?)(?:\n|$)',
            'modalidade': r'modalidade\s*:?\s*(.+?)(?:\n|$)'
        }

        import re

        for field_name, pattern in field_patterns.items():
            match = re.search(pattern, full_text.lower())
            if match:
                form_fields.append(FormField(
                    field_name=field_name,
                    field_value=match.group(1).strip(),
                    field_type='text',
                    confidence=0.9,
                    position={'x': 0, 'y': 0, 'width': 0, 'height': 0}
                ))

        return form_fields

    def _extract_from_pdf(
        self,
        pdf_content: bytes,
        filename: str,
        extract_tables: bool,
        detect_layout: bool,
        extract_forms: bool
    ) -> OCRResult:
        """Extrai texto de PDF (com ou sem OCR)."""
        if not PDF_SUPPORT:
            return self._pdf_fallback(pdf_content, filename)

        try:
            # Tentar extração direta (PDF com texto)
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())

            full_text = '\n\n'.join(text_parts)

            # Se texto extraído é muito pequeno, tentar OCR
            if len(full_text.strip()) < 100 and self.use_vision_api:
                self.logger.info("📸 PDF parece escaneado, usando OCR...")
                return self._extract_pdf_with_ocr(pdf_content, filename, extract_tables, detect_layout, extract_forms)

            return OCRResult(
                text=full_text,
                confidence=0.95,
                language='pt',
                tables=[],
                layout_blocks=[],
                form_fields=[],
                metadata={'file_type': 'pdf', 'method': 'direct_extraction'},
                processing_time=0,
                method='pypdf'
            )

        except Exception as e:
            self.logger.warning(f"⚠️  PDF direct extraction failed: {e}, trying OCR...")
            if self.use_vision_api:
                return self._extract_pdf_with_ocr(pdf_content, filename, extract_tables, detect_layout, extract_forms)
            else:
                return self._pdf_fallback(pdf_content, filename)

    def _extract_pdf_with_ocr(
        self,
        pdf_content: bytes,
        filename: str,
        extract_tables: bool,
        detect_layout: bool,
        extract_forms: bool
    ) -> OCRResult:
        """Extrai texto de PDF usando OCR (para PDFs escaneados)."""
        try:
            # Converter PDF para imagens
            images = convert_from_bytes(pdf_content)

            all_text = []
            all_tables = []
            all_layout_blocks = []
            all_form_fields = []

            # Processar cada página
            for page_num, image in enumerate(images):
                self.logger.info(f"🔍 Processing page {page_num + 1}/{len(images)}")

                # Converter PIL Image para bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()

                # Extrair com Vision API
                page_result = self._extract_with_vision_api(
                    img_bytes, f"{filename}_page_{page_num}",
                    extract_tables, detect_layout, extract_forms
                )

                all_text.append(page_result.text)
                all_tables.extend(page_result.tables)
                all_layout_blocks.extend(page_result.layout_blocks)
                all_form_fields.extend(page_result.form_fields)

            return OCRResult(
                text='\n\n---PAGE BREAK---\n\n'.join(all_text),
                confidence=0.90,
                language='pt',
                tables=all_tables,
                layout_blocks=all_layout_blocks,
                form_fields=all_form_fields,
                metadata={
                    'file_type': 'pdf',
                    'method': 'ocr',
                    'pages': len(images)
                },
                processing_time=0,
                method='vision_api_pdf'
            )

        except Exception as e:
            self.logger.error(f"❌ PDF OCR failed: {e}")
            return self._pdf_fallback(pdf_content, filename)

    def _extract_from_docx(self, docx_content: bytes, filename: str) -> OCRResult:
        """Extrai texto de documento Word (.docx)."""
        if not DOCX_SUPPORT:
            return self._docx_fallback(docx_content, filename)

        try:
            doc = DocxDocument(io.BytesIO(docx_content))

            # Extrair texto de parágrafos
            paragraphs = [p.text for p in doc.paragraphs]
            full_text = '\n\n'.join(paragraphs)

            # Extrair tabelas
            tables = []
            for table in doc.tables:
                cells = []
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cells.append(TableCell(
                            row=row_idx,
                            col=col_idx,
                            text=cell.text,
                            confidence=1.0
                        ))

                if cells:
                    tables.append(ExtractedTable(
                        rows=len(table.rows),
                        cols=len(table.columns),
                        cells=cells,
                        confidence=1.0,
                        position={'x': 0, 'y': 0, 'width': 0, 'height': 0}
                    ))

            return OCRResult(
                text=full_text,
                confidence=1.0,
                language='pt',
                tables=tables,
                layout_blocks=[],
                form_fields=[],
                metadata={'file_type': 'docx'},
                processing_time=0,
                method='docx'
            )

        except Exception as e:
            self.logger.error(f"❌ DOCX extraction failed: {e}")
            return self._docx_fallback(docx_content, filename)

    def _extract_image_fallback(self, image_content: bytes, filename: str) -> OCRResult:
        """Fallback para extração de imagem sem Vision API."""
        self.stats['fallback_calls'] += 1

        return OCRResult(
            text="[OCR não disponível - Google Vision API não configurada]",
            confidence=0.0,
            language='pt',
            tables=[],
            layout_blocks=[],
            form_fields=[],
            metadata={'file_type': 'image', 'fallback': True},
            processing_time=0,
            method='fallback'
        )

    def _pdf_fallback(self, pdf_content: bytes, filename: str) -> OCRResult:
        """Fallback para PDF sem bibliotecas disponíveis."""
        self.stats['fallback_calls'] += 1

        return OCRResult(
            text="[Extração de PDF não disponível - PyPDF2 não instalado]",
            confidence=0.0,
            language='pt',
            tables=[],
            layout_blocks=[],
            form_fields=[],
            metadata={'file_type': 'pdf', 'fallback': True},
            processing_time=0,
            method='fallback'
        )

    def _docx_fallback(self, docx_content: bytes, filename: str) -> OCRResult:
        """Fallback para DOCX sem python-docx."""
        self.stats['fallback_calls'] += 1

        return OCRResult(
            text="[Extração de DOCX não disponível - python-docx não instalado]",
            confidence=0.0,
            language='pt',
            tables=[],
            layout_blocks=[],
            form_fields=[],
            metadata={'file_type': 'docx', 'fallback': True},
            processing_time=0,
            method='fallback'
        )

    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas de uso do serviço."""
        return {
            **self.stats,
            'vision_api_enabled': self.use_vision_api,
            'success_rate': (
                (self.stats['total_extractions'] - self.stats['errors']) /
                self.stats['total_extractions'] * 100
            ) if self.stats['total_extractions'] > 0 else 0
        }
